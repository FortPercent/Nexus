import httpx, jwt, io, os
from datetime import datetime, timedelta, timezone
import openpyxl, zipfile

SECRET=os.getenv('OPENWEBUI_JWT_SECRET','6WYGSa8e7EBsSeG3')
USER='ce1d405b-0b5c-4faf-8864-010e2611b900'
tok = jwt.encode({'id':USER,'exp':datetime.now(timezone.utc)+timedelta(hours=1)},SECRET,algorithm='HS256')
H = {'Authorization':f'Bearer {tok}'}
BASE = 'http://localhost:8000/admin/api'

fails = 0
def run(name, fn):
    global fails
    try:
        note = fn() or ''
        print(f'[PASS] {name} {note}')
    except AssertionError as e:
        print(f'[FAIL] {name} - {e}'); fails += 1
    except Exception as e:
        print(f'[FAIL] {name} - {type(e).__name__}: {e}'); fails += 1

def upload(filename, content, mime='application/octet-stream'):
    files = {'file': (filename, content, mime)}
    r = httpx.post(f'{BASE}/personal/files', headers=H, files=files, timeout=60)
    if r.status_code != 200:
        raise AssertionError(f'status={r.status_code} body={r.text[:200]}')
    return r.json()

def list_personal():
    r = httpx.get(f'{BASE}/personal/files', headers=H, timeout=10)
    return [f['name'] for f in r.json()]

def delete_test_files():
    rows = httpx.get(f'{BASE}/personal/files', headers=H, timeout=10).json()
    for f in rows:
        if f['name'].startswith('TEST_'):
            httpx.delete(f"{BASE}/personal/files/{f['id']}", headers=H, timeout=10)

# 清旧
delete_test_files()

def t_xlsx():
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = 'Sheet1'
    ws.append(['项目','模型','状态'])
    ws.append(['AI Infra','Qwen3.5-122B','稳定'])
    ws.append(['AI Infra Cache','DeepSeek-V3','测试中'])
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    r = upload('TEST_sheet.xlsx', buf.getvalue(), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    assert 'TEST_sheet.xlsx' in (r.get('uploaded') or []), f'expected display name without .md, got {r}'
    files = list_personal()
    assert 'TEST_sheet.xlsx.md' in files, f'not in list: {files}'
    return f'→ {r["uploaded"]}'

def t_csv():
    content = '项目,负责人\nAI Infra,wuxn5\nCache,jinyx5'.encode('utf-8')
    r = upload('TEST_team.csv', content, 'text/csv')
    assert 'TEST_team.csv' in (r.get('uploaded') or [])
    return f'→ {r["uploaded"]}'

def t_zip():
    # 造 zip：内含 csv + md
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w') as z:
        z.writestr('data.csv', 'a,b\n1,2')
        z.writestr('notes.md', '# 注释\n内容')
    r = upload('TEST_pack.zip', buf.getvalue(), 'application/zip')
    uploaded = r.get('uploaded') or []
    # 预期：data.csv.md + notes.md 被展开 (路径前缀为 TEST_pack__)
    has_csv = any('data.csv' in n for n in uploaded)
    has_md = any('notes.md' in n for n in uploaded)
    assert has_csv and has_md, f'missing expected: {uploaded}'
    return f'→ {uploaded}'

def t_docx():
    from docx import Document
    d = Document()
    d.add_heading('测试标题', 0)
    d.add_paragraph('段落内容测试 123')
    d.add_heading('子标题', 1)
    d.add_paragraph('列表项 A', style='List Bullet')
    t = d.add_table(rows=2, cols=2)
    t.cell(0,0).text='字段'; t.cell(0,1).text='值'
    t.cell(1,0).text='环境'; t.cell(1,1).text='生产'
    buf = io.BytesIO(); d.save(buf); buf.seek(0)
    r = upload('TEST_doc.docx', buf.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    assert 'TEST_doc.docx' in (r.get('uploaded') or []), f'expected display name without .md: {r}'
    return f'→ {r["uploaded"]}'

def t_pdf_passthrough():
    # 最小 pdf header
    content = b'%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\ntrailer<</Size 1/Root 1 0 R>>\n%%EOF'
    r = upload('TEST_pass.pdf', content, 'application/pdf')
    assert 'TEST_pass.pdf' in (r.get('uploaded') or [])
    return f'→ {r["uploaded"]}'

def t_unsupported_ext():
    try:
        upload('TEST_x.exe', b'noop')
        raise AssertionError('应 400')
    except AssertionError as e:
        if '400' in str(e):
            return 'rejected 400'
        raise

def t_png_rejected_with_clear_msg():
    """04-22 加: Letta 后端 415 拒图. adapter 必须早 400 + 提示文案, 不再静默吞 200+uploaded=[].
    回归点: PASSTHROUGH_EXTS 不许包含 png/jpg/jpeg."""
    # 最小 PNG: 8 字节 signature + 一个空 IHDR, 不需要真合法图, file_processor 看扩展名就拒
    png_bytes = b'\x89PNG\r\n\x1a\n' + b'\x00' * 16
    try:
        upload('TEST_screenshot.png', png_bytes, 'image/png')
        raise AssertionError('应 400')
    except AssertionError as e:
        msg = str(e)
        if '400' not in msg:
            raise
        # 验提示文案命中预期关键词, 不只是 status 对
        if '图片' not in msg and 'image' not in msg.lower():
            raise AssertionError(f'400 但提示文案不提"图片": {msg[:200]}')
        return 'rejected 400 with image hint'

run('xlsx → markdown', t_xlsx)
run('csv → markdown', t_csv)
run('docx → markdown', t_docx)
run('zip 递归展开', t_zip)
run('pdf 原样透传', t_pdf_passthrough)
run('未知扩展 → 400', t_unsupported_ext)
run('png → 400 + 图片提示', t_png_rejected_with_clear_msg)

delete_test_files()
print()
print(f'==== {7-fails}/7 ====')
import sys; sys.exit(fails)
